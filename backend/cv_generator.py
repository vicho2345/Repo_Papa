from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand colors
NAVY = RGBColor(0x1F, 0x3A, 0x5C)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x88, 0x88, 0x88)


def _add_rule(doc):
    """Thin blue line under section headers."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_header(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    _add_rule(doc)


def generate_cv_docx(adapted_data: dict) -> bytes:
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Default style ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ════════════════════════════════════════
    # HEADER
    # ════════════════════════════════════════
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(2)
    r = p_name.add_run("RONALD SARABIA MELGAREJO")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(6)
    r = p_contact.add_run(
        "rasarabiam@gmail.com  |  +56 9 9535 9803  |  linkedin.com/in/ronald-sarabia-melgarejo-29a04b157"
    )
    r.font.size = Pt(9.5)
    r.font.color.rgb = GRAY

    _add_rule(doc)

    # ════════════════════════════════════════
    # PRESENTACIÓN PROFESIONAL
    # ════════════════════════════════════════
    _section_header(doc, "Presentación Profesional")

    presentacion = adapted_data.get("adapted_presentacion", "")
    if not presentacion:
        # Fallback: original presentation
        presentacion = (
            "Profesional con más de 29 años de experiencia en construcción de proyectos de "
            "infraestructura eléctrica y sector de la construcción, Constructor Civil, Diplomado "
            "en Gestión de la Construcción.\n\n"
            "Especialista senior en Construcción de Líneas de Alta Tensión y Subestaciones Eléctricas. "
            "Vasta experiencia en cargos de Jefe de Construcción, Jefe General, Jefe de Terreno y "
            "Administrador de Contratos para Codelco, Transelec, Minera Escondida, AngloAmerican, "
            "Los Pelambres, Collahuasi, Chilquinta y CGE Transmisión.\n\n"
            "Responsable, metódico, comprometido con la seguridad, calidad, medio ambiente y la "
            "mejora continua en todos los procesos."
        )

    for para_text in presentacion.split("\n\n"):
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)

    # Keywords badge (if adapted)
    keywords = adapted_data.get("keywords", [])
    if keywords:
        p_kw = doc.add_paragraph()
        p_kw.paragraph_format.space_before = Pt(2)
        p_kw.paragraph_format.space_after = Pt(0)
        r = p_kw.add_run("Keywords: " + "  ·  ".join(keywords))
        r.font.size = Pt(9)
        r.font.color.rgb = LIGHT_GRAY
        r.italic = True

    # ════════════════════════════════════════
    # EXPERIENCIA PROFESIONAL
    # ════════════════════════════════════════
    _section_header(doc, "Experiencia Profesional")

    experiences = [
        {
            "company": "ASA INGENIERIA, CONSTRUCCION Y MANTENIMIENTO SPA",
            "period": "Febrero – Junio 2026",
            "role": "Jefe de Construcción",
            "items": [
                "Proyecto Reemplazo Fibra Óptica zona 3, mandante Transelec S.A.",
            ],
        },
        {
            "company": "BBOSCH S.A. – TRANSBOSCH",
            "period": "Junio 2025 – Enero 2026",
            "role": "Administrador de Contratos",
            "items": [
                "División de Ingeniería y Construcción de sistemas eléctricos.",
            ],
        },
        {
            "company": "CONSORCIO KIPREOS INPROLEC",
            "period": "Enero 2023 – Abril 2025",
            "role": "Jefe General de Construcción",
            "items": [
                "Proyecto STE C20+ – Construcción, Precomisionamiento y Puesta en Marcha Líneas Aéreas y Subestaciones en 220 kV para Compañía Minera Doña Inés de Collahuasi.",
                "Alcances: SE Tarapacá · LAT 220kV Tarapacá–Puerto Collahuasi · SE Puerto Collahuasi (PS1) · SE Geoglifos · LAT 2x220kV Geoglifos–Caliche · SE Caliche (PS2).",
                "Logro: Equipos de alto rendimiento. Cumplimiento del 40% de avance a diciembre 2023 y término de obra en plazo (abril 2025).",
            ],
        },
        {
            "company": "BBOSCH S.A.",
            "period": "2000 – Agosto 2022",
            "role": "Administrador de Contratos / Jefe de Construcción / Jefe de Terreno / Jefe Oficina Técnica",
            "items": [
                "22 años en empresa líder en construcción de sistemas eléctricos. Participación en más de 25 proyectos a nivel nacional.",
                "",
                "Proyectos destacados – Administrador de Contratos / Jefe de Construcción (2011–2022):",
                "· 2019–2022 | Transelec – Línea 2x500kV Pichirropulli–Nueva Puerto Montt (140 km, USD 100M). Implementación Last Planner. Gestión exitosa durante estallido social y pandemia.",
                "· 2016–2018 | Codelco División Chuquicamata – Línea 220kV Encuentro–DMH–Tchitack + salas eléctricas para explotación subterránea (USD 35M). Metodología Obeya Lean.",
                "· 2015 | Codelco – Reconstrucción fast track Línea 110kV Llanta 2, 21 estructuras destruidas por aluvión Atacama (USD 4M).",
                "· 2013–2015 | Codelco VP – Línea distribución 110kV con acceso vía andinistas y helicóptero (USD 25M).",
                "",
                "Proyectos destacados – Jefe de Terreno (2004–2011):",
                "· 2008 | AngloAmerican Chile – By Pass Línea 66kV interior Mina Los Bronces (EPC).",
                "· 2007 | Electroandina – Línea 1x220kV SE Laberinto–SE Gaby, Antofagasta (EPC llave en mano).",
                "· 2004 | Minera Escondida – Línea 13,8kV Premina + Línea 69kV + SE Principal Escondida Norte.",
            ],
        },
        {
            "company": "SECTOR EDIFICACIÓN",
            "period": "1995 – 2000",
            "role": "Jefe de Obra / Jefe de Terreno",
            "items": [
                "Constructoras Romeral Ltda. · Jorge Rojas Figueroa · Juan Carlos Ruiz S.A. · Nualart y Cía. Ltda. · Inmobiliaria Cerro Maulén Ltda.",
            ],
        },
    ]

    for exp in experiences:
        # Company + period
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(0)
        r_co = p.add_run(exp["company"])
        r_co.bold = True
        r_co.font.size = Pt(10.5)
        r_per = p.add_run(f"  ·  {exp['period']}")
        r_per.font.size = Pt(9.5)
        r_per.font.color.rgb = GRAY

        # Role
        p_role = doc.add_paragraph()
        p_role.paragraph_format.space_before = Pt(0)
        p_role.paragraph_format.space_after = Pt(2)
        r_role = p_role.add_run(exp["role"])
        r_role.italic = True
        r_role.font.size = Pt(10)
        r_role.font.color.rgb = BLUE

        # Items
        for item in exp["items"]:
            if item == "":
                p_blank = doc.add_paragraph()
                p_blank.paragraph_format.space_after = Pt(1)
                continue
            p_item = doc.add_paragraph()
            p_item.paragraph_format.space_before = Pt(0)
            p_item.paragraph_format.space_after = Pt(1)
            p_item.paragraph_format.left_indent = Inches(0.15)
            r_item = p_item.add_run(item)
            r_item.font.size = Pt(10)

    # ════════════════════════════════════════
    # FORMACIÓN ACADÉMICA
    # ════════════════════════════════════════
    _section_header(doc, "Formación Académica")

    for degree, institution, year in [
        ("Constructor Civil", "Universidad de la Frontera", "1996"),
        ("Diplomado en Gestión de la Construcción", "Universidad Católica de Chile", "2013"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        r_d = p.add_run(f"{degree}  –  ")
        r_d.bold = True
        r_d.font.size = Pt(10.5)
        r_i = p.add_run(f"{institution}, {year}")
        r_i.font.size = Pt(10)

    p_detail = doc.add_paragraph(
        "Cursos: Administración de Proyectos · Liderazgo · Desarrollo Organizacional · Gestión y Mejoramiento de la Calidad"
    )
    p_detail.paragraph_format.space_before = Pt(0)
    p_detail.paragraph_format.space_after = Pt(2)
    p_detail.paragraph_format.left_indent = Inches(0.15)
    p_detail.runs[0].font.size = Pt(9.5)
    p_detail.runs[0].font.color.rgb = GRAY

    # ════════════════════════════════════════
    # CURSOS Y CERTIFICACIONES
    # ════════════════════════════════════════
    _section_header(doc, "Cursos y Certificaciones")

    courses = [
        "Gestión de Contratos en Proyectos de Construcción",
        "El Último Planificador (Last Planner) – Lean Construction",
        "Primavera P3  ·  MS Project  ·  Autocad R14",
        "ISO 9001-2000",
        "Prevención de Riesgos en Minería",
        "Construyendo con Cero Daño",
    ]
    for course in courses:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)
        r = p.add_run(f"• {course}")
        r.font.size = Pt(10)

    # Footer line
    job = adapted_data.get("job", {})
    if job.get("title"):
        _add_rule(doc)
        p_footer = doc.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_footer = p_footer.add_run(f"CV adaptado para: {job['title']} — {job['company']}")
        r_footer.font.size = Pt(8)
        r_footer.font.color.rgb = LIGHT_GRAY
        r_footer.italic = True

    # ── Serialize to bytes ──
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
